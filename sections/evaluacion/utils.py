"""
Utilidades para el módulo de Evaluación
Funciones auxiliares para procesamiento de datos y generación de actas
"""
import pandas as pd
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches


# ============================================================================
# FUNCIONES PARA DESEMPLEADOS - ACTAS INDIVIDUALES
# ============================================================================

def procesar_cronograma_desempleados(archivo_excel):
    """
    Procesa el archivo Excel de cronograma para desempleados
    
    Args:
        archivo_excel: Archivo Excel subido por el usuario
        
    Returns:
        dict: Diccionario con los datos extraídos del cronograma
        
    TODO: Implementar según la estructura real del Excel
    """
    try:
        df = pd.read_excel(archivo_excel)
        
        # Aquí se implementará la extracción de datos según la estructura real
        # Por ahora retornamos un ejemplo de estructura
        datos_cronograma = {
            'fecha_inicio': None,
            'fecha_fin': None,
            'horas_totales': 0,
            'modulos': [],
            'sesiones': [],
            'formadores': []
        }
        
        return datos_cronograma
        
    except Exception as e:
        raise Exception(f"Error al procesar cronograma: {str(e)}")


def procesar_asistencias_desempleados(archivo_excel):
    """
    Procesa el archivo Excel de asistencias para desempleados
    
    Args:
        archivo_excel: Archivo Excel subido por el usuario
        
    Returns:
        dict: Diccionario con los datos de asistencia procesados
        
    TODO: Implementar según la estructura real del Excel
    """
    try:
        df = pd.read_excel(archivo_excel)
        
        # Aquí se implementará la extracción de datos según la estructura real
        datos_asistencias = {
            'alumnos': [],
            'porcentajes_asistencia': {},
            'total_sesiones': 0,
            'asistencias_por_fecha': {}
        }
        
        return datos_asistencias
        
    except Exception as e:
        raise Exception(f"Error al procesar asistencias: {str(e)}")


def generar_acta_individual_desempleados(plantilla_word, datos_cronograma, datos_asistencias):
    """
    Genera las actas individuales para cada alumno
    
    Args:
        plantilla_word: Documento Word plantilla
        datos_cronograma: Datos extraídos del cronograma
        datos_asistencias: Datos extraídos de las asistencias
        
    Returns:
        list: Lista de documentos Word generados (uno por alumno)
        
    TODO: Implementar según los campos de la plantilla
    """
    try:
        documentos_generados = []
        
        # Por cada alumno, generar un acta individual
        # TODO: Implementar la lógica de generación
        
        return documentos_generados
        
    except Exception as e:
        raise Exception(f"Error al generar actas individuales: {str(e)}")


# ============================================================================
# FUNCIONES PARA DESEMPLEADOS - ACTAS TRANSVERSALES
# ============================================================================

def generar_acta_transversal_desempleados(plantilla_word, datos_cronograma, datos_asistencias):
    """
    Genera el acta transversal con información de todo el grupo
    
    Args:
        plantilla_word: Documento Word plantilla
        datos_cronograma: Datos extraídos del cronograma
        datos_asistencias: Datos extraídos de las asistencias
        
    Returns:
        Document: Documento Word generado
        
    TODO: Implementar según los campos de la plantilla
    """
    try:
        # TODO: Implementar la lógica de generación del acta transversal
        doc = Document(plantilla_word)
        
        return doc
        
    except Exception as e:
        raise Exception(f"Error al generar acta transversal: {str(e)}")


# ============================================================================
# FUNCIONES AUXILIARES GENERALES
# ============================================================================

def calcular_porcentaje_asistencia(asistencias_alumno, total_sesiones):
    """
    Calcula el porcentaje de asistencia de un alumno
    
    Args:
        asistencias_alumno: Lista o conteo de asistencias del alumno
        total_sesiones: Total de sesiones del curso
        
    Returns:
        float: Porcentaje de asistencia
    """
    if total_sesiones == 0:
        return 0.0
    
    return (asistencias_alumno / total_sesiones) * 100


def formatear_fecha(fecha, formato="%d/%m/%Y"):
    """
    Formatea una fecha según el formato especificado
    
    Args:
        fecha: Objeto datetime o string
        formato: Formato de salida
        
    Returns:
        str: Fecha formateada
    """
    if isinstance(fecha, str):
        return fecha
    
    try:
        return fecha.strftime(formato)
    except:
        return str(fecha)


def reemplazar_marcadores_word(doc, datos):
    """
    Reemplaza marcadores en un documento Word con los datos proporcionados
    
    Args:
        doc: Documento Word
        datos: Diccionario con los datos para reemplazar
        
    Returns:
        Document: Documento con los marcadores reemplazados
    """
    # Reemplazar en párrafos
    for paragraph in doc.paragraphs:
        for key, value in datos.items():
            marcador = f"{{{{{key}}}}}"
            if marcador in paragraph.text:
                paragraph.text = paragraph.text.replace(marcador, str(value))
    
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in datos.items():
                    marcador = f"{{{{{key}}}}}"
                    if marcador in cell.text:
                        cell.text = cell.text.replace(marcador, str(value))
    
    return doc


def validar_estructura_excel(df, columnas_requeridas):
    """
    Valida que un DataFrame tenga las columnas requeridas
    
    Args:
        df: DataFrame de pandas
        columnas_requeridas: Lista de nombres de columnas requeridas
        
    Returns:
        tuple: (bool, list) - (es_valido, columnas_faltantes)
    """
    columnas_df = set(df.columns)
    columnas_req = set(columnas_requeridas)
    
    columnas_faltantes = columnas_req - columnas_df
    
    return len(columnas_faltantes) == 0, list(columnas_faltantes)


def exportar_documento_a_bytes(doc):
    """
    Exporta un documento Word a bytes para descarga
    
    Args:
        doc: Documento Word
        
    Returns:
        bytes: Documento en formato bytes
    """
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()