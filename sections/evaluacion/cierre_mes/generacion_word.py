import os
# Añadir poppler al PATH del proceso de Python
poppler_path = r"C:\Users\Arancha\Desktop\Arancha\poppler-24.02.0\Library\bin"
os.environ["PATH"] = poppler_path + os.pathsep + os.environ["PATH"]


from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def construir_observaciones(ayudas, dias_aula, dias_empresa, justificantes, dias_lectivos, faltas):
    """Construye el texto de observaciones para un alumno"""
    observaciones = []
    
    if ayudas:
        for ayuda in ayudas:
            if isinstance(ayuda, dict):
                tipo = ayuda.get('tipo', '')
                if tipo:
                    observaciones.append(tipo)
            elif isinstance(ayuda, str):
                observaciones.append(ayuda)
    
    if dias_empresa > 0:
        observaciones.append(f"Empresa: {dias_empresa} días")
    
    if justificantes > 0:
        observaciones.append(f"{justificantes} faltas justificadas")
    
    return " / ".join(observaciones) if observaciones else ""


def formatear_nombre_alumno(nombre_completo):
    """
    Convierte "APELLIDO1 APELLIDO2, NOMBRE" a "NOMBRE APELLIDO1 APELLIDO2"
    
    Ejemplos:
        "GARCÍA PÉREZ, MARÍA" → "MARÍA GARCÍA PÉREZ"
        "LÓPEZ, JUAN" → "JUAN LÓPEZ"
    """
    if ',' in nombre_completo:
        partes = nombre_completo.split(',')
        apellidos = partes[0].strip()
        nombre = partes[1].strip() if len(partes) > 1 else ''
        return f"{nombre} {apellidos}".strip()
    else:
        # Si no tiene coma, se asume que ya está en el formato correcto
        return nombre_completo.strip()


def generar_parte_mensual(template_path, output_path, datos_documento):
    """
    Genera el parte mensual SOLO rellenando la tabla de alumnos
    La parte superior queda como está en el template (SIN valores por defecto)
    """
    
    try:
        print("\n" + "="*80)
        print("GENERACIÓN DEL PARTE MENSUAL - SOLO TABLA ALUMNOS")
        print("="*80)
        
        print(f"\n📂 Cargando template: {template_path}")
        doc = Document(template_path)
        
        if len(doc.tables) == 0:
            print("❌ No hay tablas")
            return False
        
        tabla = doc.tables[0]
        print(f"✅ Tabla: {len(tabla.rows)} filas x {len(tabla.columns)} columnas")
        
        # Extraer datos (SIN valores por defecto)
        alumnos = datos_documento.get('alumnos', [])
        
        print(f"\n📋 Datos:")
        print(f"  Alumnos a procesar: {len(alumnos)}")
        print(f"\n⚠️  NOTA: La cabecera del documento NO se modifica automáticamente")
        print(f"         Debes rellenar manualmente: expediente, mes, días lectivos")
        
        # Buscar fila de encabezados
        fila_inicio_alumnos = None
        for i in range(10):
            for j in range(len(tabla.rows[i].cells)):
                if 'Nombre y Apellidos' in tabla.rows[i].cells[j].text:
                    fila_inicio_alumnos = i + 1
                    print(f"\n✅ Fila de datos de alumnos: {fila_inicio_alumnos}")
                    break
            if fila_inicio_alumnos:
                break
        
        if not fila_inicio_alumnos:
            print("❌ No se encontró la fila de encabezados")
            return False
        
        # Buscar columnas
        fila_encabezado = tabla.rows[fila_inicio_alumnos - 1]
        
        col_nombre = None
        col_nif = None
        col_faltas = None
        col_obs = None
        
        for j, celda_enc in enumerate(fila_encabezado.cells):
            texto_enc = celda_enc.text.strip()
            if 'Nombre y Apellidos' in texto_enc:
                col_nombre = j
                print(f"✅ Columna Nombre: {j}")
            elif 'NIF' in texto_enc:
                col_nif = j
                print(f"✅ Columna NIF: {j}")
            elif 'faltas' in texto_enc.lower() and 'Observ' not in texto_enc:
                col_faltas = j
                print(f"✅ Columna Faltas: {j}")
            elif 'Observaciones' in texto_enc:
                col_obs = j
                print(f"✅ Columna Observaciones: {j}")
        
        # Verificar que se encontraron las columnas
        if col_nombre is None or col_nif is None or col_faltas is None or col_obs is None:
            print("❌ No se encontraron todas las columnas necesarias")
            return False
        
        # RELLENAR ALUMNOS
        print(f"\n👥 Rellenando tabla de alumnos...")
        print("-" * 100)
        
        alumnos_procesados = 0
        
        for idx, alumno in enumerate(alumnos[:20]):  # Máximo 20 alumnos
            fila_idx = fila_inicio_alumnos + idx
            if fila_idx >= len(tabla.rows):
                print(f"⚠️  Se alcanzó el límite de filas de la tabla")
                break
            
            fila = tabla.rows[fila_idx]
            
            try:
                # Obtener datos del alumno
                nombre_original = alumno.get('nombre', '')
                dni = alumno.get('dni', alumno.get('nif', ''))
                faltas = alumno.get('faltas', 0)
                observaciones = alumno.get('observaciones', '')
                
                # Formatear nombre: de "APELLIDO1 APELLIDO2, NOMBRE" a "NOMBRE APELLIDO1 APELLIDO2"
                nombre_formateado = formatear_nombre_alumno(nombre_original)
                
                # Escribir NOMBRE
                celda = fila.cells[col_nombre]
                celda.text = nombre_formateado
                for p in celda.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
                
                # Escribir DNI/NIF
                celda = fila.cells[col_nif]
                celda.text = dni
                for p in celda.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(9)
                
                # Escribir FALTAS
                celda = fila.cells[col_faltas]
                celda.text = str(faltas)
                for p in celda.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(9)
                
                # Escribir OBSERVACIONES
                celda = fila.cells[col_obs]
                celda.text = observaciones
                for p in celda.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)
                
                # Log
                nombre_log = nombre_formateado[:40]
                obs_log = observaciones[:35] if observaciones else '-'
                print(f"{idx+1:2}. {nombre_log:40} | {dni:12} | Faltas: {faltas:2} | {obs_log}")
                
                alumnos_procesados += 1
                
            except Exception as e:
                print(f"❌ Error con alumno {idx+1}: {e}")
                import traceback
                traceback.print_exc()
        
        print("-" * 100)
        print(f"✅ {alumnos_procesados} alumnos procesados correctamente")
        
        # Guardar
        print(f"\n💾 Guardando documento...")
        doc.save(output_path)
        
        if os.path.exists(output_path):
            tamanio = os.path.getsize(output_path)
            if tamanio > 10000:
                print(f"✅ Documento guardado: {output_path}")
                print(f"📊 Tamaño: {tamanio:,} bytes")
                return True
            else:
                print(f"❌ Archivo muy pequeño ({tamanio} bytes)")
                return False
        else:
            print("❌ No se pudo guardar")
            return False
        
    except Exception as e:
        print(f"\n❌ ERROR GENERAL: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    # Prueba con formato real
    datos = {
        'alumnos': [
            {
                'nombre': 'GARCÍA PÉREZ, MARÍA',  # Se convertirá a: MARÍA GARCÍA PÉREZ
                'dni': '12345678A',
                'faltas': 2,
                'observaciones': 'Beca de transporte'
            },
            {
                'nombre': 'LÓPEZ SÁNCHEZ, JUAN CARLOS',  # Se convertirá a: JUAN CARLOS LÓPEZ SÁNCHEZ
                'dni': '87654321B',
                'faltas': 0,
                'observaciones': 'Empresa: 4 días'
            },
        ]
    }
    
    template = r"C:\Users\Arancha\Desktop\Arancha\Repos\sections\evaluacion\cierre_mes\template_original.docx"
    output = r"C:\Users\Arancha\Desktop\Arancha\Repos\PRUEBA_FORMATO_NOMBRE.docx"
    
    generar_parte_mensual(template, output, datos)