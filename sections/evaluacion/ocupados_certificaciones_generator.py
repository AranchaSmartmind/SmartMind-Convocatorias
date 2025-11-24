"""
Generador de Certificados Word para Ocupados
Rellena la plantilla oficial con los datos de cada alumno
"""

from docx import Document
from docx.shared import Pt
from typing import Dict
import io


class CertificacionOcupadosGenerator:
    """Genera certificados Word individuales"""
    
    def __init__(self, plantilla_path: str):
        """
        Inicializa el generador
        
        Args:
            plantilla_path: Ruta a la plantilla .docx
        """
        self.plantilla_path = plantilla_path
        
    def generar_certificado(self, datos: Dict) -> bytes:
        """
        Genera un certificado para un alumno
        
        Args:
            datos: Diccionario con todos los datos del alumno
            
        Returns:
            Bytes del documento Word generado
        """
        doc = Document(self.plantilla_path)
        
        tabla = doc.tables[0]
        
        self._rellenar_celda(tabla, 6, 3, "")

        self._rellenar_celda(tabla, 6, 17, datos['expediente'])

        self._rellenar_celda(tabla, 8, 3, datos['director'])

        self._rellenar_celda(tabla, 9, 7, datos['centro'])

        self._rellenar_celda(tabla, 10, 6, datos['codigo_centro'])

        self._rellenar_celda(tabla, 15, 5, datos['nombre_alumno'])

        self._rellenar_celda(tabla, 15, 19, datos['dni_alumno'])
        
        fila_18 = tabla.rows[17]
        celda_modulo = fila_18.cells[0]

        for p in celda_modulo.paragraphs:
            p.clear()

        if celda_modulo.paragraphs:
            p = celda_modulo.paragraphs[0]
        else:
            p = celda_modulo.add_paragraph()
        
        run = p.add_run(datos['nombre_modulo'])
        run.font.size = Pt(9)

        self._rellenar_celda(tabla, 18, 1, datos['codigo_modulo'])

        self._rellenar_celda(tabla, 18, 11, datos['nivel'])

        self._rellenar_celda(tabla, 18, 16, datos['horas'])

        self._rellenar_celda(tabla, 19, 2, datos['fecha_inicio'])

        self._rellenar_celda(tabla, 19, 10, datos['fecha_fin'])

        self._rellenar_celda(tabla, 22, 0, datos['codigo_modulo'])
        
        fila_23 = tabla.rows[22]
        for idx in range(6, 19):
            try:
                celda = fila_23.cells[idx]

                for paragraph in celda.paragraphs:
                    paragraph.clear()
                
                if not celda.paragraphs:
                    celda.add_paragraph()
                
                paragraph = celda.paragraphs[0]
                run = paragraph.add_run(datos['nombre_modulo'])
                run.font.size = Pt(9)
                break
            except:
                pass

        self._rellenar_celda(tabla, 22, 19, datos['horas'])

        self._rellenar_celda(tabla, 22, 21, datos['calificacion'])

        fila_ciudad = tabla.rows[24]
        for celda in fila_ciudad.cells:
            if 'En' in celda.text and len(celda.text.strip()) <= 3:
                for paragraph in celda.paragraphs:
                    for run in paragraph.runs:
                        if 'En' in run.text:
                            run.text = f"En {datos['ciudad']}"
                            break
                break

        fila_firma = tabla.rows[28]
        for celda in fila_firma.cells:
            if 'Fdo.:' in celda.text or 'Fdo.' in celda.text:
                for paragraph in celda.paragraphs:
                    for run in paragraph.runs:
                        if 'Fdo.' in run.text:
                            run.text = f"Fdo.: {datos['director']}"
                            break
                break

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output.read()
    
    def _rellenar_celda(self, tabla, fila_idx: int, celda_idx: int, valor: str):
        """
        Rellena una celda específica de la tabla
        
        Args:
            tabla: Tabla del documento
            fila_idx: Índice de fila (0-based)
            celda_idx: Índice de celda (0-based)
            valor: Valor a insertar
        """
        try:
            celda = tabla.rows[fila_idx].cells[celda_idx]
            
            for paragraph in celda.paragraphs:
                for run in paragraph.runs:
                    run.text = ""

            if not celda.paragraphs:
                celda.add_paragraph()
            
            celda.paragraphs[0].text = str(valor)
            
        except Exception as e:
            print(f" Error rellenando celda [{fila_idx}, {celda_idx}]: {e}")
    
    def generar_nombre_archivo(self, datos: Dict) -> str:
        """
        Genera nombre de archivo para el certificado
        Formato: NOMBRE_APELLIDO1_APELLIDO2_DNI
        
        Args:
            datos: Datos del alumno
            
        Returns:
            Nombre de archivo (sin extensión)
        """
        nombre_completo = datos['nombre_alumno']
        partes = nombre_completo.split()
        
        if len(partes) >= 3:
            nombre = partes[0]
            apellido1 = partes[1]
            apellido2 = partes[2]
            nombre_formateado = f"{nombre}_{apellido1}_{apellido2}"
        elif len(partes) == 2:
            nombre = partes[0]
            apellido1 = partes[1]
            nombre_formateado = f"{nombre}_{apellido1}"
        else:
            nombre_formateado = nombre_completo.replace(' ', '_')
        
        dni = datos['dni_alumno']
        
        return f"{nombre_formateado}_{dni}".upper()


def generar_certificado_ocupado(plantilla_path: str, datos: Dict) -> tuple:
    """
    Función principal para generar un certificado
    
    Args:
        plantilla_path: Ruta a la plantilla
        datos: Datos del alumno
        
    Returns:
        Tuple (bytes del documento, nombre de archivo)
    """
    generator = CertificacionOcupadosGenerator(plantilla_path)
    certificado_bytes = generator.generar_certificado(datos)
    nombre_archivo = generator.generar_nombre_archivo(datos)
    
    return certificado_bytes, nombre_archivo