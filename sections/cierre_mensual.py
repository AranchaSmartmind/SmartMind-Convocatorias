"""
SERVICIO CIERRE MENSUAL - Procesador dinámico
Extrae automáticamente toda la información de los archivos
"""

import streamlit as st
import pandas as pd
import pdfplumber
import re
from datetime import datetime
import os
from docx import Document
from docx.shared import Pt
import tempfile

class ProcesadorCierreMensual:
    def __init__(self):
        self.mes_actual = None
        self.ano_actual = None
        
    def procesar_completo(self, archivos, plantilla_path):
        """Procesa todos los archivos y genera el parte mensual"""
        try:
            st.info(" Detectando información del curso...")
            self.mes_actual, self.ano_actual = self._detectar_mes_ano(archivos)
            
            st.info(" Extrayendo información del curso...")
            curso_info = self._extraer_info_curso(archivos)
            
            st.info(" Extrayendo lista de alumnos...")
            alumnos = self._extraer_alumnos(archivos)
            
            st.info(" Procesando asistencia...")
            asistencia = self._extraer_asistencia(archivos, alumnos)
            
            st.info(" Extrayendo información de ayudas...")
            ayudas = self._extraer_ayudas(archivos)
            
            st.info(" Procesando justificantes...")
            justificantes = self._extraer_justificantes(archivos)
            
            st.info(" Calculando faltas...")
            faltas = self._calcular_faltas(asistencia, justificantes)
            
            st.info(" Generando observaciones...")
            observaciones = self._generar_observaciones(alumnos, ayudas, faltas, justificantes, curso_info)
            
            st.info(" Generando documento final...")
            output_path = self._generar_documento(plantilla_path, curso_info, alumnos, faltas, observaciones)
            
            return True, output_path, {
                'total_alumnos': len(alumnos),
                'alumnos_con_faltas': sum(1 for f in faltas.values() if f > 0),
                'curso': curso_info.get('nombre', 'Curso detectado'),
                'codigo': curso_info.get('codigo', ''),
                'centro': curso_info.get('centro', 'INTERPROS NEXT GENERATION SLU'),
                'dias_lectivos': curso_info.get('dias_lectivos', '2'),
                'mes': self.mes_actual
            }
            
        except Exception as e:
            return False, None, {'error': str(e)}
    
    def _detectar_mes_ano(self, archivos):
        """Detecta mes y año automáticamente"""
        for nombre, archivo in archivos.items():
            if archivo and archivo.name.endswith('.pdf'):
                try:
                    with pdfplumber.open(archivo) as pdf:
                        texto = pdf.pages[0].extract_text()
                        fecha_match = re.search(r'Semana del:\s*(\d{2}/\d{2}/\d{4})', texto)
                        if fecha_match:
                            fecha = datetime.strptime(fecha_match.group(1), '%d/%m/%Y')
                            return fecha.strftime('%B').upper(), str(fecha.year)
                except:
                    pass
        
        # Por defecto, mes y año actual
        ahora = datetime.now()
        return ahora.strftime('%B').upper(), str(ahora.year)
    
    def _extraer_info_curso(self, archivos):
        """Extrae información del curso dinámicamente"""
        info = {
            'nombre': 'OPERACIONES AUXILIARES DE SERVICIOS ADMINISTRATIVOS Y GENERALES',
            'codigo': '', 
            'centro': 'INTERPROS NEXT GENERATION SLU', 
            'dias_lectivos': '2'
        }
        
        for nombre, archivo in archivos.items():
            if not archivo:
                continue
                
            try:
                if archivo.name.endswith('.pdf'):
                    with pdfplumber.open(archivo) as pdf:
                        texto = pdf.pages[0].extract_text()
                        
                        # Buscar código
                        if not info['codigo']:
                            codigo_match = re.search(r'(\d{4}/\d{3,4})', texto)
                            if codigo_match:
                                info['codigo'] = codigo_match.group(1)
                        
                        # Buscar días lectivos
                        if not info['dias_lectivos']:
                            dias_match = re.search(r'Días lectivos:\s*(\d+)', texto)
                            if dias_match:
                                info['dias_lectivos'] = dias_match.group(1)
                
                elif archivo.name.endswith('.xlsx'):
                    df = pd.read_excel(archivo, nrows=10)
                    # Buscar en celdas
                    for cell in df.values.flatten():
                        if pd.notna(cell):
                            cell_str = str(cell)
                            if re.search(r'\d{4}/\d+', cell_str) and not info['codigo']:
                                info['codigo'] = re.search(r'\d{4}/\d+', cell_str).group(0)
                                
            except Exception as e:
                st.warning(f" Error leyendo {archivo.name}: {e}")
        
        return info
    
    def _extraer_alumnos(self, archivos):
        """Extrae lista de alumnos"""
        alumnos = []
        
        # Buscar en Excel
        if archivos.get('excel'):
            try:
                df = pd.read_excel(archivos['excel'])
                for idx, row in df.iterrows():
                    nombre = str(row.get('APELLIDOS, NOMBRE', '')).strip()
                    dni = str(row.get('DNI', '')).strip()
                    
                    if nombre and dni and nombre != 'nan' and dni != 'nan' and len(dni) >= 8:
                        alumnos.append({
                            'nombre_original': nombre,
                            'nombre_formateado': self._formatear_nombre(nombre),
                            'nif': dni
                        })
            except Exception as e:
                st.error(f" Error leyendo Excel: {e}")
        
        # Buscar en PDF de becas
        if archivos.get('becas'):
            try:
                with pdfplumber.open(archivos['becas']) as pdf:
                    texto = pdf.pages[0].extract_text()
                    lineas = texto.split('\n')
                    
                    for linea in lineas:
                        match = re.match(r'^(\d+)\s+([A-ZÁÉÍÓÚÑ\s,]+?)\s+([A-Z0-9]{8,9})\s', linea)
                        if match:
                            nombre = match.group(2).strip()
                            nif = match.group(3).strip()
                            
                            # Evitar duplicados
                            if not any(a['nif'] == nif for a in alumnos):
                                alumnos.append({
                                    'nombre_original': nombre,
                                    'nombre_formateado': self._formatear_nombre(nombre),
                                    'nif': nif
                                })
            except Exception as e:
                st.error(f" Error leyendo PDF becas: {e}")
        
        # Ordenar alfabéticamente
        alumnos.sort(key=lambda x: x['nombre_formateado'])
        return alumnos
    
    def _formatear_nombre(self, nombre):
        """Formatea nombre APELLIDOS, NOMBRE → NOMBRE APELLIDOS"""
        if ',' in nombre:
            partes = nombre.split(',')
            return f"{partes[1].strip()} {partes[0].strip()}".strip()
        return nombre
    
    def _extraer_asistencia(self, archivos, alumnos):
        """Extrae asistencia de partes de firma"""
        asistencia = {}
        
        if archivos.get('parte_centro'):
            try:
                with pdfplumber.open(archivos['parte_centro']) as pdf:
                    texto = pdf.pages[0].extract_text()
                    nifs_presentes = re.findall(r'[0-9XYZ][0-9]{7}[A-Z]', texto)
                    
                    # Marcar presentes (0 faltas)
                    for nif in nifs_presentes:
                        asistencia[nif] = 0
                    
                    # Marcar faltantes (2 faltas para julio)
                    for alumno in alumnos:
                        if alumno['nif'] not in asistencia:
                            asistencia[alumno['nif']] = 2
                            
            except Exception as e:
                st.error(f" Error leyendo parte centro: {e}")
        
        return asistencia
    
    def _extraer_ayudas(self, archivos):
        """Extrae información de ayudas y becas"""
        ayudas = {}
        
        if archivos.get('becas'):
            try:
                with pdfplumber.open(archivos['becas']) as pdf:
                    texto = pdf.pages[0].extract_text()
                    lineas = texto.split('\n')
                    
                    for linea in lineas:
                        match = re.match(r'^(\d+)\s+([A-ZÁÉÍÓÚÑ\s,]+?)\s+([A-Z0-9]{8,9})\s+(.+)$', linea)
                        if match:
                            nif = match.group(3)
                            resto = match.group(4)
                            
                            ayudas[nif] = {
                                'transporte': 'X' in resto,
                                'conciliacion': resto.count('X') >= 2,
                                'beca': 'DISCAPACIDAD' if 'DISCAPACIDAD' in resto.upper() else None
                            }
            except Exception as e:
                st.error(f" Error extrayendo ayudas: {e}")
        
        return ayudas
    
    def _extraer_justificantes(self, archivos):
        """Extrae justificantes automáticamente"""
        justificantes = {}
        
        if archivos.get('justificantes'):
            try:
                with pdfplumber.open(archivos['justificantes']) as pdf:
                    for pagina in pdf.pages:
                        texto = pagina.extract_text()
                        if texto:
                            nifs = re.findall(r'[0-9XYZ][0-9]{7}[A-Z]', texto)
                            for nif in nifs:
                                if nif in justificantes:
                                    justificantes[nif] += 1
                                else:
                                    justificantes[nif] = 1
            except Exception as e:
                st.error(f" Error leyendo justificantes: {e}")
        
        return justificantes
    
    def _calcular_faltas(self, asistencia, justificantes):
        """Calcula faltas finales considerando justificantes"""
        faltas_finales = {}
        
        for nif, faltas_iniciales in asistencia.items():
            justificadas = justificantes.get(nif, 0)
            faltas_finales[nif] = max(0, faltas_iniciales - justificadas)
        
        return faltas_finales
    
    def _generar_observaciones(self, alumnos, ayudas, faltas, justificantes, curso_info):
        """Genera observaciones automáticamente"""
        observaciones = {}
        
        for alumno in alumnos:
            nif = alumno['nif']
            partes = []
            
            # Ayudas
            ayuda_info = ayudas.get(nif, {})
            if ayuda_info.get('transporte'):
                partes.append('Transporte')
            if ayuda_info.get('conciliacion'):
                partes.append('Conciliación')
            if ayuda_info.get('beca'):
                partes.append(f"Beca {ayuda_info['beca']}")
            
            # Días asistidos
            faltas_alumno = faltas.get(nif, 0)
            dias_asistidos = int(curso_info.get('dias_lectivos', 2)) - faltas_alumno
            if partes:  # Solo mostrar días si tiene ayudas
                partes.append(f"{dias_asistidos}d")
            
            # Justificantes
            justificadas = justificantes.get(nif, 0)
            if justificadas > 0:
                partes.append(f"{justificadas} justificada{'s' if justificadas > 1 else ''}")
            
            observaciones[nif] = " + ".join(partes) if partes else ""
        
        return observaciones
    
    def _generar_documento(self, plantilla_path, curso_info, alumnos, faltas, observaciones):
        """Genera el documento Word final"""
        try:
            doc = Document(plantilla_path)
            tabla = doc.tables[0]
            
            # Rellenar información del curso
            tabla.rows[0].cells[10].text = curso_info.get('codigo', '')
            tabla.rows[1].cells[1].text = curso_info.get('nombre', '')
            tabla.rows[2].cells[3].text = curso_info.get('centro', '')
            tabla.rows[3].cells[3].text = self.mes_actual
            tabla.rows[3].cells[10].text = str(curso_info.get('dias_lectivos', '2'))
            
            # Rellenar alumnos
            fila_inicio = 6
            for i, alumno in enumerate(alumnos[:20]):  # Máximo 20 alumnos
                fila_idx = fila_inicio + i
                if fila_idx < len(tabla.rows):
                    fila = tabla.rows[fila_idx]
                    nif = alumno['nif']
                    
                    fila.cells[1].text = str(i + 1)
                    fila.cells[2].text = alumno['nombre_formateado']
                    fila.cells[4].text = nif
                    fila.cells[8].text = str(faltas.get(nif, 0))
                    fila.cells[9].text = observaciones.get(nif, '')
                    
                    # Ajustar fuentes
                    for cell_idx in [2, 4, 8, 9]:
                        for paragraph in fila.cells[cell_idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8 if cell_idx != 9 else 7)
            
            # Guardar en archivo temporal
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            return temp_file.name
            
        except Exception as e:
            st.error(f" Error generando documento: {e}")
            return None