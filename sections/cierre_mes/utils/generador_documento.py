"""
Módulo de generación de documentos Word
Rellena el template con los datos procesados
"""

from docx import Document

def construir_observaciones(ayudas, dias_aula, dias_empresa, justificantes, dias_lectivos):
    """
    Construye el texto de observaciones según formato oficial
    
    Ejemplos:
    - Transporte+ conciliación: 23
    - Discapacidad+23
    - 1 falta justificada
    
    Args:
        ayudas: Lista de ayudas del alumno
        dias_aula: Número de días de aula
        dias_empresa: Número de días de empresa
        justificantes: Número de justificantes
        dias_lectivos: Total de días lectivos
    
    Returns:
        str: Texto de observaciones
    """
    partes = []
    
    # Ayudas
    if ayudas:
        texto_ayudas = '+ '.join(ayudas)
        if 'Discapacidad' in ayudas and len(ayudas) == 1:
            partes.append(f"Discapacidad+{dias_lectivos}")
        else:
            partes.append(f"{texto_ayudas}: {dias_lectivos}")
    
    # Días de aula
    if dias_aula > 0:
        partes.append(f"{dias_aula} días aula")
    
    # Justificantes
    if justificantes > 0:
        if justificantes == 1:
            partes.append("1 falta justificada")
        else:
            partes.append(f"{justificantes} faltas justificadas")
    
    return '\n'.join(partes)

def rellenar_template(template_path, datos):
    """
    Rellena el template Word con los datos del parte mensual
    
    Args:
        template_path: Ruta al template Word
        datos: Diccionario con todos los datos
    
    Returns:
        Document: Documento Word rellenado
    """
    try:
        doc = Document(template_path)
        
        if not doc.tables:
            print("Error: No se encontró tabla en el template")
            return None
        
        table = doc.tables[0]
        
        # FILA 1: Número de curso
        table.rows[0].cells[10].text = datos['numero_curso']
        
        # FILA 2: Especialidad
        for i in range(1, 11):
            table.rows[1].cells[i].text = datos['especialidad']
        
        # FILA 3: Centro
        centro = datos.get('centro', 'INTERPROS NEXT GENERATION SLU')
        for i in range(3, 11):
            table.rows[2].cells[i].text = centro
        
        # FILA 4: Mes y Días lectivos
        mes_texto = f"Mes de {datos['mes']}"
        table.rows[3].cells[1].text = mes_texto
        table.rows[3].cells[2].text = mes_texto
        
        dias_texto = f"Número de días lectivos: {datos['dias_lectivos']}"
        for i in range(5, 10):
            table.rows[3].cells[i].text = dias_texto
        
        # FILAS 7-26: Alumnos
        for idx, alumno in enumerate(datos['alumnos']):
            if idx >= 20:  # Máximo 20 alumnos
                break
            
            fila_idx = 6 + idx
            row = table.rows[fila_idx]
            
            # Número
            row.cells[1].text = str(idx + 1)
            
            # Nombre y Apellidos
            row.cells[2].text = alumno['nombre']
            row.cells[3].text = alumno['nombre']
            
            # NIF
            for i in range(4, 8):
                row.cells[i].text = alumno['dni']
            
            # Número de faltas
            row.cells[8].text = str(alumno['faltas'])
            
            # Observaciones
            observaciones = alumno['observaciones']
            row.cells[9].text = observaciones
            row.cells[10].text = observaciones
        
        return doc
    
    except Exception as e:
        print(f"Error rellenando template: {e}")
        return None