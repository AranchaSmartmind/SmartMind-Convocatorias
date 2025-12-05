"""
Sección de Memorias - Generación automática con Claude
INCLUYE TABLAS de aspectos mejor/peor valorados
"""
import streamlit as st
import os
import tempfile
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# PROMPT OPTIMIZADO para replicar estilo del documento original
PROMPT_TEMPLATE = """Analiza los siguientes resultados de las encuestas de satisfacción del curso {numero_curso}.

Los resultados se basan en las medias obtenidas de las respuestas de los alumnos en una escala Likert 1-4, donde un valor más alto es más positivo.

DATOS DE LAS ENCUESTAS:
{datos_encuestas}

{comentarios_seccion}

ESTRUCTURA REQUERIDA:

PÁRRAFO INTRODUCTORIO
Comienza con un párrafo que contextualice el informe.

ANÁLISIS GENERAL: ASPECTOS MEJOR Y PEOR VALORADOS
Identifica los 3-5 aspectos mejor valorados y los 3-5 peor valorados del curso.

VALORACIÓN DETALLADA POR ASPECTOS
Analiza cada uno de estos aspectos:
1. Organización del Curso
2. Contenido y Metodología
3. Duración y Horario
4. Formadores / Tutores
5. Medios Didácticos
6. Instalaciones y medios técnicos
7. Guías y materiales didácticos
8. Mecanismos para la evaluación del aprendizaje
9. Valoración General Curso
10. Grado Satisfacción General Curso

Para cada aspecto:
- Indica la valoración general del área
- Menciona los puntos específicos con sus medias
- Usa datos numéricos concretos

COMENTARIOS Y PROPUESTAS DE MEJORA
{comentarios_instruccion}

FORMATO CRÍTICO (muy importante):
- NO uses markdown (sin **, sin *, sin #, sin ---, sin bullets)
- Escribe solo en TEXTO PLANO
- Para resaltar información importante, escribe: [NEGRITA]texto importante[/NEGRITA]
- Ejemplo: "El área de Organización del Curso ha obtenido la media más alta del informe."
  Siguiente línea: "[NEGRITA]La valoración general de esta área es de 3.70.[/NEGRITA]"
- Los títulos de sección también van con [NEGRITA]título[/NEGRITA]
- Párrafos fluidos y profesionales
- Menciona siempre las medias numéricas específicas"""


def extraer_datos_excel_memorias(excel_path):
    """
    Extrae datos del Excel de encuestas de satisfacción
    VERSIÓN MEJORADA: detecta automáticamente la estructura
    """
    try:
        # Leer Excel sin asumir primera fila como header
        df = pd.read_excel(excel_path, header=None)
        
        print("\nForma del Excel:", df.shape)
        
        datos = {
            'preguntas': [],
            'comentarios': []
        }
        
        # Buscar fila de encabezados
        header_row = None
        for idx, row in df.iterrows():
            row_str = ' '.join([str(x) for x in row if pd.notna(x)]).upper()
            if 'PREGUNTA' in row_str or 'MEDIA' in row_str:
                header_row = idx
                break
        
        if header_row is not None:
            df = pd.read_excel(excel_path, header=header_row)
        else:
            df = pd.read_excel(excel_path)
        
        # Buscar columnas
        col_pregunta = None
        col_media = None
        
        for col in df.columns:
            col_str = str(col).upper()
            if 'PREGUNTA' in col_str:
                col_pregunta = col
            elif 'MEDIA' in col_str:
                col_media = col
        
        if not col_pregunta:
            col_pregunta = df.columns[0]
        
        if not col_media:
            col_media = df.columns[1] if len(df.columns) > 1 else None
        
        # Procesar filas
        for idx, row in df.iterrows():
            pregunta = str(row.get(col_pregunta, '')).strip()
            
            if not pregunta or pregunta == 'nan' or len(pregunta) < 3:
                continue
            
            if col_media and pd.notna(row.get(col_media)):
                try:
                    media = float(row.get(col_media, 0))
                    
                    item = {
                        'pregunta': pregunta,
                        'media': media,
                        'desv_std': 0,
                        'resp_1': 0,
                        'resp_2': 0,
                        'resp_3': 0,
                        'resp_4': 0
                    }
                    
                    # Extraer datos adicionales si existen
                    if len(df.columns) > 2:
                        try:
                            item['desv_std'] = float(row.iloc[2]) if pd.notna(row.iloc[2]) else 0
                        except:
                            pass
                    
                    if len(df.columns) > 3:
                        try:
                            item['resp_1'] = int(row.iloc[3]) if pd.notna(row.iloc[3]) else 0
                        except:
                            pass
                    
                    if len(df.columns) > 4:
                        try:
                            item['resp_2'] = int(row.iloc[4]) if pd.notna(row.iloc[4]) else 0
                        except:
                            pass
                    
                    if len(df.columns) > 5:
                        try:
                            item['resp_3'] = int(row.iloc[5]) if pd.notna(row.iloc[5]) else 0
                        except:
                            pass
                    
                    if len(df.columns) > 6:
                        try:
                            item['resp_4'] = int(row.iloc[6]) if pd.notna(row.iloc[6]) else 0
                        except:
                            pass
                    
                    datos['preguntas'].append(item)
                    
                except:
                    if len(pregunta) > 10:
                        datos['comentarios'].append(pregunta)
            else:
                if len(pregunta) > 10:
                    datos['comentarios'].append(pregunta)
        
        print(f"Preguntas extraídas: {len(datos['preguntas'])}")
        print(f"Comentarios extraídos: {len(datos['comentarios'])}")
        
        return datos
        
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        return None


def construir_prompt_memorias(datos_excel, numero_curso):
    """Construye el prompt para Gemini"""
    
    datos_texto = ""
    for i, item in enumerate(datos_excel['preguntas'], 1):
        datos_texto += f"""
Pregunta {i}: {item['pregunta']}
- Media: {item['media']:.2f}
- Desviación estándar: {item['desv_std']:.2f}
- Respuestas: 1={item['resp_1']}, 2={item['resp_2']}, 3={item['resp_3']}, 4={item['resp_4']}
"""
    
    if datos_excel['comentarios']:
        comentarios_seccion = f"""
COMENTARIOS DE LOS ALUMNOS:
{chr(10).join([f"- {com}" for com in datos_excel['comentarios']])}
"""
        comentarios_instruccion = "Analiza estos comentarios y propón mejoras específicas."
    else:
        comentarios_seccion = ""
        comentarios_instruccion = "No hay comentarios registrados."
    
    prompt = PROMPT_TEMPLATE.format(
        numero_curso=numero_curso,
        datos_encuestas=datos_texto,
        comentarios_seccion=comentarios_seccion,
        comentarios_instruccion=comentarios_instruccion
    )
    
    return prompt


def generar_memoria_con_gemini(prompt, api_key):
    """Genera memoria con Google Gemini"""
    try:
        import requests
        import json
        
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={api_key}"
        
        headers = {'Content-Type': 'application/json'}
        
        data = {
            "contents": [{
                "parts": [{
                    "text": prompt
                }]
            }],
            "generationConfig": {
                "temperature": 0.7,
                "maxOutputTokens": 16000
            }
        }
        
        print("Llamando a Gemini 2.5 Flash...")
        response = requests.post(url, headers=headers, json=data, timeout=120)
        
        if response.status_code == 200:
            result = response.json()
            
            print("Respuesta recibida, parseando...")
            print(f"Estructura: {result.keys()}")
            
            # Intentar diferentes estructuras de respuesta
            if 'candidates' in result and len(result['candidates']) > 0:
                candidate = result['candidates'][0]
                
                # Opción 1: candidates[0]['content']['parts'][0]['text']
                if 'content' in candidate:
                    content = candidate['content']
                    if 'parts' in content and len(content['parts']) > 0:
                        texto = content['parts'][0]['text']
                        print(f"✓ Respuesta recibida: {len(texto)} caracteres")
                        return texto
                
                # Opción 2: candidates[0]['text']
                if 'text' in candidate:
                    texto = candidate['text']
                    print(f"✓ Respuesta recibida: {len(texto)} caracteres")
                    return texto
                
                # Opción 3: candidates[0]['output']
                if 'output' in candidate:
                    texto = candidate['output']
                    print(f"✓ Respuesta recibida: {len(texto)} caracteres")
                    return texto
            
            # Si ninguna estructura funcionó
            print("Error: No se pudo extraer el texto")
            print(f"Estructura completa: {json.dumps(result, indent=2)[:500]}")
            return None
        else:
            print(f"✗ Error API: {response.status_code}")
            print(f"Respuesta: {response.text}")
            return None
        
    except Exception as e:
        print(f"✗ Error Gemini API: {e}")
        import traceback
        traceback.print_exc()
        return None


def insertar_tabla_aspectos(doc, datos_excel, tipo='mejor'):
    """Inserta tabla con aspectos mejor o peor valorados"""
    
    # Ordenar preguntas por media
    preguntas_ordenadas = sorted(datos_excel['preguntas'], key=lambda x: x['media'], reverse=(tipo=='mejor'))
    
    # Tomar top 5 (o menos si hay menos preguntas)
    num_items = min(5, len(preguntas_ordenadas))
    
    if tipo == 'mejor':
        top_preguntas = preguntas_ordenadas[:num_items]
    else:
        top_preguntas = list(reversed(preguntas_ordenadas[-num_items:]))
    
    if not top_preguntas:
        return
    
    # Calcular media promedio
    media_promedio = sum([p['media'] for p in top_preguntas]) / len(top_preguntas)
    
    # Título de la tabla
    titulo = f"Aspectos {'Mejor' if tipo == 'mejor' else 'Peor'} Valorados (Puntuación media - {media_promedio:.2f})"
    
    # Crear tabla
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # CRÍTICO: Configurar tabla para que NO se divida entre páginas
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Evitar división de tabla
    cantSplit = OxmlElement('w:cantSplit')
    tblPr.append(cantSplit)
    
    # Encabezado
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = titulo
    hdr_cells[1].text = 'Puntuación'
    
    # Formato encabezado (fondo amarillo, negrita)
    for cell in hdr_cells:
        cell_pr = cell._element.get_or_add_tcPr()
        cell_shading = OxmlElement('w:shd')
        cell_shading.set(qn('w:fill'), 'FFFF00')  # Amarillo
        cell_pr.append(cell_shading)
        
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    
    # Añadir filas con datos
    for pregunta in top_preguntas:
        row_cells = table.add_row().cells
        
        # CRÍTICO: Configurar la fila para que NO se divida
        tr = table.rows[-1]._element
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)
        
        # Texto de la pregunta
        texto = pregunta['pregunta']
        
        # Buscar palabras clave para poner en negrita
        palabras_clave = [
            'número de alumnos', 'bien organizado', 'conocen los temas',
            'comprensibles y adecuados', 'instalaciones han sido apropiadas',
            'medios técnicos', 'ampliado conocimientos', 'favorecido el desarrollo',
            'combinación adecuada', 'mejorado mis posibilidades', 'duración del curso',
            'han sido adecuados'
        ]
        
        # Añadir texto con formato
        p = row_cells[0].paragraphs[0]
        p.clear()
        
        texto_lower = texto.lower()
        ultima_pos = 0
        partes_encontradas = []
        
        for palabra in palabras_clave:
            pos = texto_lower.find(palabra.lower())
            if pos != -1 and pos >= ultima_pos:
                partes_encontradas.append((pos, pos + len(palabra), palabra))
        
        # Ordenar por posición
        partes_encontradas.sort()
        
        ultima_pos = 0
        for inicio, fin, palabra in partes_encontradas:
            # Texto antes
            if inicio > ultima_pos:
                p.add_run(texto[ultima_pos:inicio])
            # Texto en negrita
            run = p.add_run(texto[inicio:fin])
            run.bold = True
            ultima_pos = fin
        
        # Resto del texto
        if ultima_pos < len(texto):
            p.add_run(texto[ultima_pos:])
        
        # Puntuación
        row_cells[1].text = f"{pregunta['media']:.2f}"
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Ajustar ancho de columnas
    table.columns[0].width = Inches(5.5)
    table.columns[1].width = Inches(1.0)
    
    # Espacio después de la tabla
    doc.add_paragraph()


def convertir_texto_a_word(texto, output_path, numero_curso, datos_excel=None):
    """Convierte texto a Word REPLICANDO ESTILO ORIGINAL"""
    try:
        doc = Document()
        
        # Título principal centrado y en negrita
        titulo = doc.add_heading(f'Informe de Análisis de Resultados - Curso {numero_curso}', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Limpiar markdown por si acaso
        texto = texto.replace('**', '')
        texto = texto.replace('###', '')
        texto = texto.replace('##', '')
        texto = texto.replace('#', '')
        texto = texto.replace('---', '')
        texto = texto.replace('***', '')
        
        # Variables para control de tablas
        ultima_linea_era_mejor = False
        ultima_linea_era_peor = False
        ya_inserto_tabla_mejor = False
        ya_inserto_tabla_peor = False
        
        lineas = texto.split('\n')
        
        for linea in lineas:
            linea = linea.strip()
            
            if not linea:
                continue
            
            # Detectar menciones de tablas
            linea_upper = linea.upper()
            if 'MEJOR VALORAD' in linea_upper and 'ASPECT' in linea_upper:
                ultima_linea_era_mejor = True
            elif 'PEOR VALORAD' in linea_upper and 'ASPECT' in linea_upper:
                ultima_linea_era_peor = True
            
            # Procesar línea con posibles [NEGRITA]
            if '[NEGRITA]' in linea:
                # Parsear y crear párrafo con negritas
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # CRÍTICO: Si es un título o subtítulo, mantenerlo con el siguiente párrafo
                # Detectar si es título (empieza con [NEGRITA] y es corto/tiene números)
                es_titulo = (
                    linea.startswith('[NEGRITA]') and 
                    (len(linea) < 150 or 
                     any(linea.startswith(f'[NEGRITA]{i}.') for i in range(1, 11)) or
                     'ANÁLISIS' in linea.upper() or
                     'VALORACIÓN' in linea.upper() or
                     'COMENTARIOS' in linea.upper())
                )
                
                if es_titulo:
                    p.paragraph_format.keep_with_next = True
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(6)
                
                # Dividir por las etiquetas
                partes = linea.split('[NEGRITA]')
                
                for i, parte in enumerate(partes):
                    if '[/NEGRITA]' in parte:
                        # Parte con negrita
                        texto_negrita, texto_normal = parte.split('[/NEGRITA]', 1)
                        
                        run_negrita = p.add_run(texto_negrita)
                        run_negrita.bold = True
                        run_negrita.font.name = 'Arial'
                        run_negrita.font.size = Pt(11)
                        
                        if texto_normal.strip():
                            run_normal = p.add_run(texto_normal)
                            run_normal.font.name = 'Arial'
                            run_normal.font.size = Pt(11)
                    else:
                        # Parte sin negrita
                        if parte.strip():
                            run = p.add_run(parte)
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
                
                # Insertar tabla si corresponde
                if ultima_linea_era_mejor and not ya_inserto_tabla_mejor and datos_excel:
                    # Mantener el párrafo anterior con la tabla
                    p.paragraph_format.keep_with_next = True
                    insertar_tabla_aspectos(doc, datos_excel, tipo='mejor')
                    ya_inserto_tabla_mejor = True
                    ultima_linea_era_mejor = False
                
                if ultima_linea_era_peor and not ya_inserto_tabla_peor and datos_excel:
                    # Mantener el párrafo anterior con la tabla
                    p.paragraph_format.keep_with_next = True
                    insertar_tabla_aspectos(doc, datos_excel, tipo='peor')
                    ya_inserto_tabla_peor = True
                    ultima_linea_era_peor = False
            
            else:
                # Línea sin etiquetas de negrita - párrafo normal
                p = doc.add_paragraph(linea)
                p.style.font.name = 'Arial'
                p.style.font.size = Pt(11)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error Word: {e}")
        import traceback
        traceback.print_exc()
        return False


def render_memorias():
    """Interfaz principal"""
    st.title("Generación de Memorias")
    st.markdown("---")
    
    tab1, tab2, tab3 = st.tabs(["Subir Excel", "Generar Memoria", "Descargar"])
    
    with tab1:
        st.subheader("Subir datos de encuestas")
        
        st.info("Excel con columnas: Pregunta, Media, Desviación Estándar, Respuesta 1-4")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            excel_file = st.file_uploader(
                "Excel con resultados",
                type=['xlsx', 'xls'],
                key="mem_excel"
            )
        
        with col2:
            numero_curso = st.text_input(
                "Número de curso",
                value="2024/XXX",
                key="mem_num_curso"
            )
        
        if excel_file:
            st.success("Excel cargado")
            with st.expander("Vista previa"):
                df = pd.read_excel(excel_file)
                st.dataframe(df.head(10))
        else:
            st.warning("Sube un archivo Excel")
    
    with tab2:
        if not excel_file:
            st.warning("Primero sube el Excel")
        else:
            st.info("Gemini de Google es GRATIS. Obtén tu API key en: https://aistudio.google.com/app/apikey")
            
            api_key = st.text_input(
                "API Key de Google Gemini (gratis)",
                type="password",
                help="Obtén tu API key gratis en https://aistudio.google.com/app/apikey",
                key="mem_api_key"
            )
            
            if not api_key:
                st.warning("Introduce tu API key de Gemini para continuar")
            
            with st.expander("Ver prompt que se usará", expanded=False):
                temp_path = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx').name
                with open(temp_path, 'wb') as f:
                    f.write(excel_file.getbuffer())
                
                datos = extraer_datos_excel_memorias(temp_path)
                if datos:
                    prompt_preview = construir_prompt_memorias(datos, numero_curso)
                    
                    st.write(f"**Prompt generado:** {len(prompt_preview)} caracteres")
                    st.write("")
                    
                    # USAR st.code() que siempre es visible
                    st.code(prompt_preview, language="text")
            
            if st.button("Generar Memoria", type="primary", use_container_width=True, disabled=not api_key):
                with st.spinner("Generando..."):
                    progress = st.progress(0)
                    status = st.empty()
                    
                    try:
                        temp_dir = tempfile.mkdtemp()
                        
                        status.text("Guardando Excel...")
                        progress.progress(10)
                        
                        excel_path = os.path.join(temp_dir, excel_file.name)
                        with open(excel_path, 'wb') as f:
                            f.write(excel_file.getbuffer())
                        
                        status.text("Extrayendo datos...")
                        progress.progress(20)
                        
                        datos = extraer_datos_excel_memorias(excel_path)
                        
                        if not datos:
                            st.error("No se pudieron extraer datos")
                            return
                        
                        status.text("Construyendo prompt...")
                        progress.progress(30)
                        
                        prompt = construir_prompt_memorias(datos, numero_curso)
                        
                        status.text("Llamando a Gemini...")
                        progress.progress(40)
                        
                        contenido = generar_memoria_con_gemini(prompt, api_key)
                        
                        if not contenido:
                            st.error("Error con Gemini API")
                            st.warning("Usa el MODO MANUAL abajo")
                            
                            st.markdown("---")
                            st.subheader("MODO MANUAL")
                            st.info("1. Copia el prompt de abajo\n2. Pégalo en https://gemini.google.com\n3. Copia la respuesta de Gemini\n4. Pégala en el campo de abajo")
                            
                            # Prompt para copiar con mejor estilo
                            st.markdown("**Paso 1: Copia este prompt**")
                            st.text_area(
                                "Prompt (selecciona todo con Ctrl+A y copia con Ctrl+C)",
                                prompt,
                                height=250,
                                key="prompt_manual",
                                help="Copia todo este texto y pégalo en https://gemini.google.com"
                            )
                            
                            st.markdown("**Paso 2: Pega aquí la respuesta de Gemini**")
                            contenido_manual = st.text_area(
                                "Respuesta de Gemini",
                                height=300,
                                key="respuesta_manual",
                                help="Copia la respuesta completa de Gemini y pégala aquí",
                                placeholder="Pega aquí la respuesta completa que te dio Gemini..."
                            )
                            
                            if contenido_manual and st.button("Generar Word con respuesta manual"):
                                status.text("Creando Word...")
                                output_path = os.path.join(temp_dir, f'Memoria_{numero_curso.replace("/", "_")}.docx')
                                exito = convertir_texto_a_word(contenido_manual, output_path, numero_curso, datos)
                                
                                if exito and os.path.exists(output_path):
                                    with open(output_path, 'rb') as f:
                                        st.session_state['mem_documento'] = f.read()
                                    st.session_state['mem_nombre'] = f'Memoria_{numero_curso.replace("/", "_")}.docx'
                                    st.success("Documento creado con respuesta manual")
                                    st.info("Ve a Descargar")
                            
                            return
                        
                        status.text("Creando Word con tablas...")
                        progress.progress(70)
                        
                        output_path = os.path.join(temp_dir, f'Memoria_{numero_curso.replace("/", "_")}.docx')
                        
                        # IMPORTANTE: pasar datos_excel para las tablas
                        exito = convertir_texto_a_word(contenido, output_path, numero_curso, datos)
                        
                        progress.progress(90)
                        
                        if exito and os.path.exists(output_path):
                            status.text("Completado")
                            progress.progress(100)
                            
                            with open(output_path, 'rb') as f:
                                st.session_state['mem_documento'] = f.read()
                            st.session_state['mem_nombre'] = f'Memoria_{numero_curso.replace("/", "_")}.docx'
                            
                            st.success("Memoria generada con tablas")
                            
                            with st.expander("Vista previa"):
                                st.text_area("Contenido", contenido, height=300)
                            
                            st.info("Ve a Descargar")
                        else:
                            st.error("Error al crear Word")
                        
                    except Exception as e:
                        st.error(f"Error: {str(e)}")
                        with st.expander("Detalles"):
                            import traceback
                            st.code(traceback.format_exc())
    
    with tab3:
        if 'mem_documento' in st.session_state:
            st.success("Documento listo")
            
            st.download_button(
                label="Descargar Memoria",
                data=st.session_state['mem_documento'],
                file_name=st.session_state['mem_nombre'],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True
            )
            
            if st.button("Nueva Memoria", use_container_width=True):
                del st.session_state['mem_documento']
                del st.session_state['mem_nombre']
                st.rerun()
        else:
            st.info("No hay documento generado")