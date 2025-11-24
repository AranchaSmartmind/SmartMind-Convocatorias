""" 
SMARTMIND- GENERADOR DE PARTE MENSUAL

Módulo para generar automáticamente el Parte Mensual a partir de:
- PDF de Otorgamiento de Becas
- Excel CTRL de Alumnos
- De momento Configuracion manual de faltas y observaciones
"""

from docx import Document
from docx.shared import Pt
import pdfplumber
import re
from pathlib import Path
from datetime import satetime, timedelta
import pandas as pd
from typing import Dict, List, Optional, Tuple
import pytesseract
from PIL import Image

def formatear_nombre(nombre_completo: str) -> str:
    """
    Convierte 'APELLIDO1 APELLIDO2, NOMBRE' a 'NOMBRE APELLIDO1 APELLIDO2'

    Args:
        nombre_completo: Nombre en formato "APELLIDOS, NOMBRE"

    Returns:
        str: Nombre formateado como "NOMBRE APELLIDO1 APELLIDO2"
    """
    
    if ',' in nombre_completo:
        partes = nombre_completo.split(',')
        apellidos = partes[0].strip()
        nombre = partes[1].strip if len(partes) > 1 else ''
        return f"{nombre} {apellidos}" .strip()
    return nombre_completo.strip()

def es_nombre_valido(nombre: str) -> bool:
    """
    Verifica si es un nombre válido (no es texto descriptivo)
    
    Args:
        nombre: Cadena a validar
    
    Returns:
        bool: True si es válido
    """
    invalidos = ['Alumnos', 'Faltan', 'SOLO', 'nan', 'None', 'que tiene']
    for inv in invalidos:
        if inv.lower() in str(nombre).lower():
            return False
    return True
    
def extraer_datos_otorgamiento(pdf_path: str) -> Dict:
    """
    Extrae información básica del PDF de otorgamiento
    
    Args:
        pdf_path: Ruta al PDF de otorgamiento
    
    Returns:
        dict: Datos extraídos (expediente, curso, días lectivos, localidad)
    """
    datos = {
        'expediente': None,
        'curso': None,
        'dias_lectivos': None,
        'localidad': None
    }
        
    try: 
        with pdfplumber.open(pdf_path) as pdf:
            texto = pdf.pages[0].extraer.text()
            
            match_expediente = re.search(r'EXPTE[:\s]+(\d{4}/\d{4}', texto)
            if match_expediente:
                datos['expediente'] = match_expediente.group(1)
                
            match_curso = re.search(r'Denominación del curso[:\s]+(.+?)(?:\n|Nº)', texto, re.IGNORECASE)
            if match_curso:
                datos['curso'] = match_curso.group(1).strtip()
                
            match_dias = re.search(r'(\d+)\s+día[s]?\s+lectivo[s]?', texto, re.IGNORECASE)
            if match_dias:
                datos['dias_lectivos'] = match_dias.group(1)
                
            match_localidad = re.search(r'Localidad[:\s]+([A-ZÁÉÍÓÚÑ\s]+)', texto)
            if match_localidad:
                datos['localidad'] = match_localidad.group(1).strip()
                
        return datos
    
    except Exception as e:
        print(f"Error al extraer datos del PDF: {e}")
        return datos
    
def extraer_ayudas_becas(pdf_otorgamiento_path: str) -> Dict:
    """
    Extrae las ayudas y becas de cada alumno del PDF de otorgamiento
    
    Returns:
        dict: {nif: {'transporte': bool, 'conciliacion': bool, 'beca': str}}
    """
    ayudas_dict = {}
    
    try:
        with pdfplumber.open(pdf_otorgamiento_path) as pdf:
            texto = pdf.pages[0].extract_text()
            lineas = texto.split('\n')
            
            for linea in lineas:
                match = re.match(r'^(\d+)\s+([A-ZÁÉÍÓÚÑ\s,]+?)\s+([A-Z0-9]{8,9})\s+(.+)$', linea)
                if match:
                    nif = match.group(3)
                    resto = match.group(4)
                    
                    num_x = resto.count('X')
                    tiene_transporte = 'X' in resto
                    tiene_conciliacion = num_x >= 2

                    beca = None
                    if 'DISCAPACIDAD' in resto.upper():
                        beca = 'DISCAPACIDAD'
                    
                    ayudas_dict[nif] = {
                        'transporte': tiene_transporte,
                        'conciliacion': tiene_conciliacion,
                        'beca': beca
                    }
        
        return ayudas_dict
        
    except Exception as e:
        print(f"Error al extraer ayudas: {e}")
        return {}
    
def extraer_alumnos_excel(excel_path: str) -> List[Dict]:
    """
    Extrae alumnos del Excel CTRL
    
    Returns:
        list: Lista de diccionarios con nombre_original y dni
    """
    alumnos = []
    
    try:
        df = pd.read_excel(excel_path)
        
        for idx, row in df.iterrows():
            nombre = str(row['APELLIDOS, NOMBRE']).strip()
            dni = str(row['DNI']).strip()
            
            if (nombre and nombre != 'nan' and 
                dni and dni != 'nan' and dni != 'None' and
                es_nombre_valido(nombre) and
                re.match(r'^[A-Z0-9]{8,9}$', dni)):
                
                alumnos.append({
                    'nombre_original': nombre,
                    'dni': dni
                })
        
        return alumnos
        
    except Exception as e:
        print(f"Error al leer Excel: {e}")
        return []
    
def extraer_alumnos_pdf_otorgamiento(pdf_path: str) -> List[Dict]:
    """
    Extrae alumnos del PDF de otorgamiento
    
    Returns:
        list: Lista de diccionarios con nombre_original y nif
    """
    alumnos = []
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            texto = pdf.pages[0].extract_text()
            lineas = texto.split('\n')
            
            for linea in lineas:
                match = re.match(r'^(\d+)\s+([A-ZÁÉÍÓÚÑ\s,]+?)\s+([A-Z]\d{7,8}[A-Z])\s', linea)
                if match:
                    nombre = match.group(2).strip()
                    nif = match.group(3).strip()
                    alumnos.append({
                        'nombre_original': nombre,
                        'nif': nif
                    })
        
        return alumnos
        
    except Exception as e:
        print(f"Error al leer PDF: {e}")
        return []
    
def unificar_alumnos(alumnos_excel: List[Dict], alumnos_pdf: List[Dict]) -> List[Dict]:
    """
    Unifica listas de alumnos eliminando duplicados por DNI/NIF
    
    Args:
        alumnos_excel: Lista de alumnos del Excel
        alumnos_pdf: Lista de alumnos del PDF
    
    Returns:
        list: Lista de alumnos únicos ordenados alfabéticamente
    """
    alumnos_dict = {}
    
    for alumno in alumnos_pdf:
        nif = alumno['nif']
        alumnos_dict[nif] = {
            'nombre_original': alumno['nombre_original'],
            'nombre_formateado': formatear_nombre(alumno['nombre_original']),
            'nif': nif
        }
    
    for alumno in alumnos_excel:
        dni = alumno['dni']
        if dni not in alumnos_dict:
            alumnos_dict[dni] = {
                'nombre_original': alumno['nombre_original'],
                'nombre_formateado': formatear_nombre(alumno['nombre_original']),
                'nif': dni
            }
    
    lista_alumnos = list(alumnos_dict.values())
    lista_alumnos.sort(key=lambda x: x['nombre_formateado'])
    
    return lista_alumnos

def aplicar_faltas_manuales(lista_alumnos: List[Dict], 
                           faltas_manuales: Dict[str, int]) -> Dict[str, int]:
    """
    Aplica las faltas manuales configuradas
    
    Args:
        lista_alumnos: Lista de alumnos
        faltas_manuales: Diccionario {nif: num_faltas}
    
    Returns:
        dict: {nif: num_faltas}
    """
    faltas_dict = {}
    
    for alumno in lista_alumnos:
        nif = alumno['nif']
        if nif in faltas_manuales:
            faltas_dict[nif] = faltas_manuales[nif]
        else:
            faltas_dict[nif] = 0
    
    return faltas_dict

def generar_observaciones(lista_alumnos: List[Dict], 
                         ayudas_dict: Dict,
                         faltas_dict: Dict[str, int],
                         dias_lectivos_totales: int,
                         observaciones_manuales: Dict[str, str],
                         faltas_justificadas: Dict[str, int]) -> Dict[str, str]:
    """
    Genera las observaciones para cada alumno
    
    Args:
        lista_alumnos: Lista de alumnos
        ayudas_dict: Diccionario con ayudas por NIF
        faltas_dict: Diccionario con faltas por NIF
        dias_lectivos_totales: Total de días lectivos
        observaciones_manuales: Observaciones predefinidas
        faltas_justificadas: Faltas justificadas por NIF
    
    Returns:
        dict: {nif: observacion}
    """
    observaciones_dict = {}
    
    for alumno in lista_alumnos:
        nif = alumno['nif']

        if observaciones_manuales and nif in observaciones_manuales:
            observacion_base = observaciones_manuales[nif]
        else:
            ayudas = ayudas_dict.get(nif, {})
            faltas = faltas_dict.get(nif, 0)
            
            dias_asistidos = dias_lectivos_totales - faltas

            partes = []
            
            if ayudas.get('transporte'):
                partes.append('Transporte')
            
            if ayudas.get('conciliacion'):
                partes.append('Conciliación')
            
            if ayudas.get('beca'):
                partes.append(f"Beca {ayudas['beca']}")
            
            if partes:
                observacion_base = ' + '.join(partes) + f": {dias_asistidos}"
            else:
                observacion_base = ""

        if faltas_justificadas and nif in faltas_justificadas:
            num_justificadas = faltas_justificadas[nif]
            if num_justificadas > 0:
                texto_justificadas = f"{num_justificadas} falta justificada" if num_justificadas == 1 else f"{num_justificadas} faltas justificadas"
                
                if observacion_base:
                    observaciones_dict[nif] = f"{observacion_base}, {texto_justificadas}"
                else:
                    observaciones_dict[nif] = texto_justificadas
            else:
                observaciones_dict[nif] = observacion_base
        else:
            observaciones_dict[nif] = observacion_base
    
    return observaciones_dict

def generar_parte_mensual(
    plantilla_path: str,
    datos_curso: Dict,
    lista_alumnos: List[Dict],
    faltas_dict: Dict[str, int],
    observaciones_dict: Dict[str, str],
    output_path: str
) -> Tuple[bool, str]:
    """
    Genera el documento Parte Mensual completo
    
    Args:
        plantilla_path: Ruta a la plantilla .docx
        datos_curso: Datos del curso (expediente, nombre, días lectivos)
        lista_alumnos: Lista de alumnos
        faltas_dict: Faltas por NIF
        observaciones_dict: Observaciones por NIF
        output_path: Ruta de salida
    
    Returns:
        tuple: (éxito, mensaje)
    """
    try:
        doc = Document(plantilla_path)
        tabla = doc.tables[0]

        tabla.rows[0].cells[10].text = datos_curso.get('expediente', '')

        tabla.rows[1].cells[1].text = datos_curso.get('curso', '')

        tabla.rows[2].cells[3].text = "INTERPROS NEXT GENERATION SLU"

        tabla.rows[3].cells[3].text = ""

        tabla.rows[3].cells[10].text = str(datos_curso.get('dias_lectivos', ''))

        fila_inicio = 6
        max_alumnos = 20
        
        for i, alumno in enumerate(lista_alumnos[:max_alumnos]):
            fila_idx = fila_inicio + i
            fila = tabla.rows[fila_idx]
            nif = alumno['nif']

            fila.cells[1].text = str(i + 1)

            celda_nombre = fila.cells[2]
            celda_nombre.text = alumno['nombre_formateado']
            for paragraph in celda_nombre.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

            celda_nif = fila.cells[4]
            celda_nif.text = nif
            for paragraph in celda_nif.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

            num_faltas = faltas_dict.get(nif, 0)
            celda_faltas = fila.cells[8]
            celda_faltas.text = str(num_faltas)
            for paragraph in celda_faltas.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

            observacion = observaciones_dict.get(nif, "")
            celda_obs = fila.cells[9]
            celda_obs.text = observacion
            for paragraph in celda_obs.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)

        doc.save(output_path)
        
        return True, f"Documento generado exitosamente: {output_path}"
        
    except Exception as e:
        return False, f"Error al generar documento: {str(e)}"


def procesar_parte_mensual(
    pdf_otorgamiento: str,
    excel_ctrl: str,
    plantilla_docx: str,
    output_path: str,
    dias_lectivos_manual: Optional[int] = None,
    faltas_manuales: Optional[Dict[str, int]] = None,
    observaciones_manuales: Optional[Dict[str, str]] = None,
    faltas_justificadas: Optional[Dict[str, int]] = None
) -> Tuple[bool, str, Dict]:
    """
    Función principal para procesar y generar el Parte Mensual
    
    Args:
        pdf_otorgamiento: Ruta al PDF de otorgamiento
        excel_ctrl: Ruta al Excel CTRL
        plantilla_docx: Ruta a la plantilla Word
        output_path: Ruta de salida
        dias_lectivos_manual: Días lectivos (opcional)
        faltas_manuales: Faltas por NIF (opcional)
        observaciones_manuales: Observaciones por NIF (opcional)
        faltas_justificadas: Faltas justificadas por NIF (opcional)
    
    Returns:
        tuple: (éxito, mensaje, estadísticas)
    """
    faltas_manuales = faltas_manuales or {}
    observaciones_manuales = observaciones_manuales or {}
    faltas_justificadas = faltas_justificadas or {}
    
    try:
        datos = extraer_datos_otorgamiento(pdf_otorgamiento)
        
        if not datos['expediente'] or not datos['curso']:
            return False, "No se pudieron extraer datos del PDF de otorgamiento", {}
        
        if dias_lectivos_manual:
            datos['dias_lectivos'] = str(dias_lectivos_manual)
        
        dias_lectivos_num = int(datos.get('dias_lectivos', 23))
        alumnos_excel = extraer_alumnos_excel(excel_ctrl)
        alumnos_pdf = extraer_alumnos_pdf_otorgamiento(pdf_otorgamiento)
        lista_alumnos = unificar_alumnos(alumnos_excel, alumnos_pdf)
        
        if not lista_alumnos:
            return False, "No se encontraron alumnos", {}
        
        faltas_dict = aplicar_faltas_manuales(lista_alumnos, faltas_manuales)

        ayudas_dict = extraer_ayudas_becas(pdf_otorgamiento)

        observaciones_dict = generar_observaciones(
            lista_alumnos, ayudas_dict, faltas_dict, dias_lectivos_num,
            observaciones_manuales, faltas_justificadas
        )

        exito, mensaje = generar_parte_mensual(
            plantilla_docx, datos, lista_alumnos, 
            faltas_dict, observaciones_dict, output_path
        )

        estadisticas = {
            'total_alumnos': len(lista_alumnos),
            'alumnos_con_faltas': sum(1 for f in faltas_dict.values() if f > 0),
            'total_faltas': sum(faltas_dict.values()),
            'alumnos_con_obs': len([o for o in observaciones_dict.values() if o]),
            'expediente': datos['expediente'],
            'curso': datos['curso'],
            'dias_lectivos': datos['dias_lectivos']
        }
        
        return exito, mensaje, estadisticas
        
    except Exception as e:
        return False, f"Error en el proceso: {str(e)}", {}































    
    
                                    
                
                
        