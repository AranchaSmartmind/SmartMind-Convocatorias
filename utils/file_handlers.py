"""
Funciones para manejar archivos Word y Excel
"""
import io
import openpyxl
import pandas as pd
import streamlit as st
from docx import Document


def rellenar_acta_desde_plantilla(plantilla_file, datos_alumnos, tipo_acta="individual", alumno_seleccionado=None):
    """Rellena una plantilla de acta Word con los datos extraídos"""
    try:
        plantilla_file.seek(0)
        doc = Document(plantilla_file)
        
        def reemplazar_en_parrafo(parrafo, buscar, reemplazar):
            if buscar in parrafo.text:
                for run in parrafo.runs:
                    if buscar in run.text:
                        run.text = run.text.replace(buscar, str(reemplazar))
        
        def reemplazar_en_tabla(tabla, buscar, reemplazar):
            for fila in tabla.rows:
                for celda in fila.cells:
                    if buscar in celda.text:
                        for parrafo in celda.paragraphs:
                            reemplazar_en_parrafo(parrafo, buscar, reemplazar)
        
        if tipo_acta == "individual" and alumno_seleccionado:
            datos = datos_alumnos.get(alumno_seleccionado, {})
            
            for parrafo in doc.paragraphs:
                reemplazar_en_parrafo(parrafo, "[NOMBRE]", alumno_seleccionado)
                reemplazar_en_parrafo(parrafo, "[DNI]", datos.get("dni", ""))
                reemplazar_en_parrafo(parrafo, "[ASISTENCIA]", datos.get("asistencia", ""))
                reemplazar_en_parrafo(parrafo, "[CALIFICACION]", datos.get("calificacion_global", ""))
                
                modulos = datos.get("modulos", {})
                for modulo, calificacion in modulos.items():
                    reemplazar_en_parrafo(parrafo, f"[{modulo}]", calificacion)
            
            for tabla in doc.tables:
                reemplazar_en_tabla(tabla, "[NOMBRE]", alumno_seleccionado)
                reemplazar_en_tabla(tabla, "[DNI]", datos.get("dni", ""))
                reemplazar_en_tabla(tabla, "[ASISTENCIA]", datos.get("asistencia", ""))
                reemplazar_en_tabla(tabla, "[CALIFICACION]", datos.get("calificacion_global", ""))
                
                modulos = datos.get("modulos", {})
                for modulo, calificacion in modulos.items():
                    reemplazar_en_tabla(tabla, f"[{modulo}]", calificacion)
        
        elif tipo_acta == "grupal":
            for tabla in doc.tables:
                if len(tabla.rows) > 2:
                    fila_encabezado = tabla.rows[0]
                    
                    for i in range(len(tabla.rows) - 1, 0, -1):
                        tabla._element.remove(tabla.rows[i]._element)
                    
                    for idx, (nombre, datos) in enumerate(datos_alumnos.items(), 1):
                        nueva_fila = tabla.add_row()
                        
                        celdas = nueva_fila.cells
                        if len(celdas) > 0:
                            celdas[0].text = str(idx)
                        if len(celdas) > 1:
                            celdas[1].text = nombre
                        if len(celdas) > 2:
                            celdas[2].text = datos.get("dni", "")
                        
                        modulos = datos.get("modulos", {})
                        col_idx = 3
                        for modulo, calificacion in modulos.items():
                            if col_idx < len(celdas):
                                celdas[col_idx].text = str(calificacion)
                                col_idx += 1
                        
                        if len(celdas) > col_idx:
                            celdas[-2].text = datos.get("asistencia", "")
                        if len(celdas) > col_idx + 1:
                            celdas[-1].text = datos.get("calificacion_global", "")
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output
    
    except Exception as e:
        st.error(f"Error al rellenar plantilla: {str(e)}")
        import traceback
        st.code(traceback.format_exc())
        return None


def visualizar_documento_word(doc_bytes):
    """Muestra el contenido de un documento Word en pantalla"""
    try:
        doc_bytes.seek(0)
        doc = Document(doc_bytes)
        
        st.markdown('<div class="custom-card">', unsafe_allow_html=True)
        
        for parrafo in doc.paragraphs:
            if parrafo.text.strip():
                if parrafo.style.name.startswith('Heading'):
                    st.markdown(f"### {parrafo.text}")
                else:
                    st.markdown(f"{parrafo.text}")
        
        for tabla in doc.tables:
            datos_tabla = []
            for fila in tabla.rows:
                datos_fila = [celda.text for celda in fila.cells]
                datos_tabla.append(datos_fila)
            
            if datos_tabla:
                df = pd.DataFrame(datos_tabla[1:], columns=datos_tabla[0])
                st.dataframe(df, use_container_width=True)
        
        st.markdown('</div>', unsafe_allow_html=True)
        
    except Exception as e:
        st.error(f"Error al visualizar documento: {str(e)}")