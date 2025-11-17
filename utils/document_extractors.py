"""
Funciones para extraer texto de diferentes tipos de documentos
"""
import re
import pandas as pd
import streamlit as st
from PIL import Image
import pytesseract
import PyPDF2
import docx


def extraer_texto_pdf(file):
    """Extrae texto de un archivo PDF"""
    try:
        if PyPDF2 is None:
            st.error("PyPDF2 no está instalado")
            return ""
        
        pdf_reader = PyPDF2.PdfReader(file)
        texto = ""
        for page in pdf_reader.pages:
            texto += page.extract_text() + "\n"
        return texto
    except Exception as e:
        st.error(f"Error al leer PDF: {str(e)}")
        return ""


def extraer_texto_imagen(file):
    """Extrae texto de una imagen usando OCR"""
    try:
        image = Image.open(file)
        texto = pytesseract.image_to_string(image, lang='spa')
        return texto
    except Exception as e:
        st.error(f"Error al procesar imagen: {str(e)}")
        return ""


def extraer_texto_word(file):
    """Extrae texto de un archivo Word"""
    try:
        doc = docx.Document(file)
        texto = ""
        for paragraph in doc.paragraphs:
            texto += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texto += cell.text + " "
            texto += "\n"
        return texto
    except Exception as e:
        st.error(f"Error al leer Word: {str(e)}")
        return ""


def extraer_texto_excel(file):
    """Extrae texto de un archivo Excel"""
    try:
        xls = pd.ExcelFile(file)
        texto = ""
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(file, sheet_name=sheet_name)
            texto += f"\n--- Hoja: {sheet_name} ---\n"
            texto += df.to_string(index=False) + "\n"
        return texto
    except Exception as e:
        st.error(f"Error al leer Excel: {str(e)}")
        return ""


def procesar_documento(file):
    """Procesa cualquier tipo de documento y extrae su contenido"""
    if file.type == "application/pdf":
        return extraer_texto_pdf(file)
    elif file.type in ["image/png", "image/jpeg", "image/jpg", "image/bmp", "image/tiff", "image/gif"]:
        return extraer_texto_imagen(file)
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extraer_texto_word(file)
    elif file.type in ["application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
        return extraer_texto_excel(file)
    else:
        st.warning(f"Tipo de archivo: {file.type}")
        try:
            return file.read().decode('utf-8')
        except:
            return ""


def extraer_datos_multiples_documentos(archivos):
    """Extrae datos de múltiples archivos de diferentes formatos"""
    datos_combinados = {
        "alumnos": {},
        "info_curso": {}
    }
    
    for archivo in archivos:
        try:
            archivo.seek(0)
            
            if archivo.type in ["application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
                try:
                    xls = pd.ExcelFile(archivo)
                    for sheet_name in xls.sheet_names:
                        df = pd.read_excel(archivo, sheet_name=sheet_name)
                        
                        for col in df.columns:
                            col_lower = str(col).lower()
                            if 'nombre' in col_lower or 'alumno' in col_lower:
                                for idx, row in df.iterrows():
                                    nombre = str(row[col]).strip().upper()
                                    if nombre and len(nombre) > 3 and not any(x in nombre for x in ['NOMBRE', 'ALUMNO', 'TOTAL']):
                                        if nombre not in datos_combinados["alumnos"]:
                                            datos_combinados["alumnos"][nombre] = {
                                                "dni": "",
                                                "modulos": {},
                                                "asistencia": "",
                                                "calificacion_global": ""
                                            }
                                        
                                        for dni_col in df.columns:
                                            if 'dni' in str(dni_col).lower():
                                                dni_val = row.get(dni_col)
                                                if pd.notna(dni_val):
                                                    datos_combinados["alumnos"][nombre]["dni"] = str(dni_val)
                                        
                                        for mod_col in df.columns:
                                            if 'MF' in str(mod_col).upper():
                                                mod_val = row.get(mod_col)
                                                if pd.notna(mod_val):
                                                    datos_combinados["alumnos"][nombre]["modulos"][str(mod_col)] = str(mod_val)
                                        
                                        for asist_col in df.columns:
                                            if '%' in str(asist_col) or 'asistencia' in str(asist_col).lower():
                                                asist_val = row.get(asist_col)
                                                if pd.notna(asist_val):
                                                    datos_combinados["alumnos"][nombre]["asistencia"] = str(asist_val)
                except Exception as e:
                    st.warning(f"Error procesando Excel {archivo.name}: {str(e)}")
            
            elif archivo.type == "application/pdf":
                texto = extraer_texto_pdf(archivo)
                if texto:
                    patron = r'(\d{8}[A-Z])\s+([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s,]+?)(?=\s*\d{8}[A-Z]|\n\n|\Z)'
                    matches = re.findall(patron, texto, re.DOTALL)
                    
                    for dni, nombre in matches:
                        nombre_limpio = nombre.strip().upper()
                        if len(nombre_limpio) > 5:
                            if nombre_limpio not in datos_combinados["alumnos"]:
                                datos_combinados["alumnos"][nombre_limpio] = {
                                    "dni": dni,
                                    "modulos": {},
                                    "asistencia": "",
                                    "calificacion_global": ""
                                }
                            else:
                                if not datos_combinados["alumnos"][nombre_limpio]["dni"]:
                                    datos_combinados["alumnos"][nombre_limpio]["dni"] = dni
            
            elif archivo.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                texto = extraer_texto_word(archivo)
                if texto:
                    patron = r'(\d{8}[A-Z])\s+([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s,]+?)(?=\s*\d{8}[A-Z]|\n\n|\Z)'
                    matches = re.findall(patron, texto, re.DOTALL)
                    
                    for dni, nombre in matches:
                        nombre_limpio = nombre.strip().upper()
                        if len(nombre_limpio) > 5:
                            if nombre_limpio not in datos_combinados["alumnos"]:
                                datos_combinados["alumnos"][nombre_limpio] = {
                                    "dni": dni,
                                    "modulos": {},
                                    "asistencia": "",
                                    "calificacion_global": ""
                                }
            
            elif archivo.type in ["image/png", "image/jpeg", "image/jpg"]:
                texto = extraer_texto_imagen(archivo)
                if texto:
                    patron = r'(\d{8}[A-Z])\s+([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s,]+?)(?=\s*\d{8}[A-Z]|\n\n|\Z)'
                    matches = re.findall(patron, texto, re.DOTALL)
                    
                    for dni, nombre in matches:
                        nombre_limpio = nombre.strip().upper()
                        if len(nombre_limpio) > 5:
                            if nombre_limpio not in datos_combinados["alumnos"]:
                                datos_combinados["alumnos"][nombre_limpio] = {
                                    "dni": dni,
                                    "modulos": {},
                                    "asistencia": "",
                                    "calificacion_global": ""
                                }
        
        except Exception as e:
            st.warning(f"Error procesando {archivo.name}: {str(e)}")
            continue
    
    return datos_combinados