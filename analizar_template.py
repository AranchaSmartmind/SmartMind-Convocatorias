from docx import Document

def analizar_template(ruta):
    """Analiza detalladamente la estructura del template"""
    doc = Document(ruta)
    
    print("=" * 80)
    print("ANÁLISIS COMPLETO DEL TEMPLATE")
    print("=" * 80)
    
    # Analizar párrafos
    print("\n📄 PÁRRAFOS:")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"Párrafo {i}: '{para.text}'")
    
    # Analizar tablas
    print(f"\n📊 NÚMERO DE TABLAS: {len(doc.tables)}")
    
    for idx_tabla, tabla in enumerate(doc.tables):
        print(f"\n{'='*60}")
        print(f"TABLA {idx_tabla + 1}")
        print(f"{'='*60}")
        print(f"Dimensiones: {len(tabla.rows)} filas x {len(tabla.columns)} columnas")
        
        # Analizar cada fila
        for i, fila in enumerate(tabla.rows):
            print(f"\n--- FILA {i} ---")
            for j, celda in enumerate(fila.cells):
                texto = celda.text.strip()
                if texto:
                    print(f"  Celda[{i},{j}]: '{texto}'")
        
        # Detectar celdas combinadas
        print("\n🔗 CELDAS COMBINADAS:")
        for i, fila in enumerate(tabla.rows):
            for j, celda in enumerate(fila.cells):
                # Verificar si es la primera aparición de la celda
                if j == 0 or celda != fila.cells[j-1]:
                    # Contar cuántas columnas ocupa
                    span = 1
                    for k in range(j+1, len(fila.cells)):
                        if fila.cells[k] == celda:
                            span += 1
                        else:
                            break
                    if span > 1:
                        print(f"  Fila {i}, Celda {j}: ocupa {span} columnas")

if __name__ == "__main__":
    ruta = r"C:\Users\Arancha\Desktop\Arancha\Repos\sections\evaluacion\cierre_mes\template_original.docx"
    analizar_template(ruta)